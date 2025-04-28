"""
RAG Components for Streamlit App
-------------------------------
This module provides simplified components for the Document Q&A Streamlit application.
It includes document processing, embedding generation, RAG system, and Q&A functionality.
"""

import os
import sys
import json
import logging
import tempfile
import hashlib
import numpy as np
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import re
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Try to import optional dependencies
try:
    import faiss
    FAISS_AVAILABLE = True
except ImportError:
    FAISS_AVAILABLE = False
    logger.warning("FAISS not available. Install with 'pip install faiss-cpu' for better vector search.")

try:
    import chromadb
    CHROMA_AVAILABLE = True
except ImportError:
    CHROMA_AVAILABLE = False
    logger.warning("ChromaDB not available. Install with 'pip install chromadb' for persistent vector storage.")

try:
    from sentence_transformers import SentenceTransformer
    HAVE_SENTENCE_TRANSFORMERS = True
except ImportError:
    HAVE_SENTENCE_TRANSFORMERS = False
    logger.warning("sentence_transformers not available. Install with 'pip install sentence-transformers'.")

try:
    import requests
    HAVE_REQUESTS = True
except ImportError:
    HAVE_REQUESTS = False
    logger.warning("requests not available. API-based embeddings will use fallback method.")

try:
    import PyPDF2
    HAVE_PYPDF2 = True
except ImportError:
    HAVE_PYPDF2 = False
    logger.warning("PyPDF2 not available. PDF processing will be limited.")

try:
    import docx
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False
    logger.warning("python-docx not available. DOCX processing will be limited.")

try:
    import csv
    HAVE_CSV = True
except ImportError:
    HAVE_CSV = False
    logger.warning("csv module not available. CSV processing will be limited.")

# -----------------------------------------------------------------------------
# Document Processing
# -----------------------------------------------------------------------------

def process_document(file_path: str, file_name: Optional[str] = None) -> Optional[Dict[str, Any]]:
    """
    Process a document file and extract content and metadata.

    Args:
        file_path: Path to the document file
        file_name: Original file name (optional)

    Returns:
        Dictionary with document content and metadata, or None if processing fails
    """
    try:
        # Use original file name if provided, otherwise use the path's basename
        if file_name is None:
            file_name = os.path.basename(file_path)

        # Get file extension
        extension = os.path.splitext(file_name)[1].lower()

        # Process based on file type
        if extension == '.txt':
            try:
                return process_text_file(file_path, file_name)
            except Exception as e:
                logger.error(f"Error processing text file {file_name}: {str(e)}")
                # Create a minimal document with error message
                return create_error_document(file_path, file_name, f"Error processing text file: {str(e)}")

        elif extension == '.pdf':
            if HAVE_PYPDF2:
                try:
                    return process_pdf_file(file_path, file_name)
                except Exception as e:
                    logger.error(f"Error processing PDF file {file_name}: {str(e)}")
                    # Create a minimal document with error message
                    return create_error_document(file_path, file_name, f"Error processing PDF file: {str(e)}")
            else:
                logger.warning(f"PDF processing requires PyPDF2. Install with 'pip install PyPDF2'")
                # Try fallback to basic text extraction
                try:
                    return process_text_file(file_path, file_name)
                except:
                    return create_error_document(file_path, file_name, "PDF processing requires PyPDF2")

        elif extension in ['.docx', '.doc']:
            if HAVE_DOCX:
                try:
                    return process_docx_file(file_path, file_name)
                except Exception as e:
                    logger.error(f"Error processing DOCX file {file_name}: {str(e)}")
                    # Create a minimal document with error message
                    return create_error_document(file_path, file_name, f"Error processing DOCX file: {str(e)}")
            else:
                logger.warning(f"DOCX processing requires python-docx. Install with 'pip install python-docx'")
                # Try fallback to basic text extraction
                try:
                    return process_text_file(file_path, file_name)
                except:
                    return create_error_document(file_path, file_name, "DOCX processing requires python-docx")

        elif extension == '.csv':
            if HAVE_CSV:
                try:
                    return process_csv_file(file_path, file_name)
                except Exception as e:
                    logger.error(f"Error processing CSV file {file_name}: {str(e)}")
                    # Create a minimal document with error message
                    return create_error_document(file_path, file_name, f"Error processing CSV file: {str(e)}")
            else:
                logger.warning(f"CSV module not available")
                # Try fallback to basic text extraction
                try:
                    return process_text_file(file_path, file_name)
                except:
                    return create_error_document(file_path, file_name, "CSV module not available")

        else:
            logger.warning(f"Unsupported file type: {extension}")
            # Try to process as text anyway
            try:
                # First check if it's a binary file
                is_binary = False
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        file.read(1024)  # Try to read as text
                except UnicodeDecodeError:
                    is_binary = True

                if is_binary:
                    logger.warning(f"File {file_name} appears to be binary and cannot be processed as text")
                    # Create a minimal document with just metadata
                    return create_error_document(file_path, file_name, "Binary file cannot be processed as text")
                else:
                    return process_text_file(file_path, file_name)
            except Exception as e:
                logger.error(f"Failed to process {file_name} as text: {str(e)}")
                return create_error_document(file_path, file_name, f"Failed to process as text: {str(e)}")

    except Exception as e:
        logger.error(f"Error processing document {file_name}: {str(e)}")
        # Create a minimal document with error message
        try:
            return create_error_document(file_path, file_name, f"Error processing document: {str(e)}")
        except:
            # Last resort fallback
            return {
                'content': f"Error processing document: {file_name}",
                'metadata': {
                    'source': file_name,
                    'file_type': 'unknown',
                    'error': str(e)
                }
            }

def create_error_document(file_path: str, file_name: str, error_message: str) -> Dict[str, Any]:
    """Create a minimal document with error information."""
    try:
        # Get basic file metadata
        file_stats = os.stat(file_path)
        file_size = file_stats.st_size
        modified_time = datetime.fromtimestamp(file_stats.st_mtime).isoformat()
    except:
        file_size = 0
        modified_time = datetime.now().isoformat()

    # Get file extension
    extension = os.path.splitext(file_name)[1].lower().lstrip('.')
    if not extension:
        extension = 'unknown'

    return {
        'content': f"[Error processing {file_name}: {error_message}]",
        'metadata': {
            'source': file_name,
            'file_type': extension,
            'file_size': file_size,
            'modified_time': modified_time,
            'error': error_message
        }
    }

def process_text_file(file_path: str, file_name: str) -> Dict[str, Any]:
    """Process a text file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()

        # Get basic file metadata
        file_stats = os.stat(file_path)
        metadata = {
            'source': file_name,
            'file_type': 'text',
            'file_size': file_stats.st_size,
            'modified_time': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
            'line_count': content.count('\n') + 1,
            'word_count': len(content.split())
        }

        # Add creation time if available (platform-dependent)
        try:
            if hasattr(file_stats, 'st_birthtime'):  # macOS
                metadata['created_time'] = datetime.fromtimestamp(file_stats.st_birthtime).isoformat()
            else:  # Fallback for other platforms
                # Just use modification time as creation time to avoid deprecated warning
                metadata['created_time'] = metadata['modified_time']
        except:
            pass

        return {
            'content': content,
            'metadata': metadata
        }
    except Exception as e:
        logger.error(f"Error processing text file {file_name}: {str(e)}")
        raise

def process_pdf_file(file_path: str, file_name: str) -> Dict[str, Any]:
    """Process a PDF file."""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)

            # Extract text content
            content = ""
            for page_num in range(len(reader.pages)):
                try:
                    page = reader.pages[page_num]
                    extracted_text = page.extract_text()
                    if extracted_text:
                        content += extracted_text + "\n\n"
                    else:
                        content += f"[Page {page_num+1} - No extractable text]\n\n"
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num+1}: {str(e)}")
                    content += f"[Page {page_num+1} - Error: {str(e)}]\n\n"

            # If no content was extracted, add a placeholder
            if not content.strip():
                content = f"[No extractable text content in {file_name}]"

            # Extract metadata
            info = reader.metadata
            metadata = {
                'source': file_name,
                'file_type': 'pdf',
                'page_count': len(reader.pages)
            }

            # Add PDF-specific metadata if available
            if info is not None:
                metadata.update({
                    'title': info.get('/Title', ''),
                    'author': info.get('/Author', ''),
                    'creator': info.get('/Creator', ''),
                    'producer': info.get('/Producer', '')
                })

            return {
                'content': content,
                'metadata': metadata
            }
    except Exception as e:
        logger.error(f"Error processing PDF file {file_name}: {str(e)}")
        raise

def process_docx_file(file_path: str, file_name: str) -> Dict[str, Any]:
    """Process a DOCX file."""
    try:
        doc = docx.Document(file_path)

        # Extract text content
        content = ""
        for para in doc.paragraphs:
            content += para.text + "\n"

        # Extract metadata
        core_properties = doc.core_properties
        metadata = {
            'source': file_name,
            'file_type': 'docx',
            'title': core_properties.title or '',
            'author': core_properties.author or '',
            'created': core_properties.created.isoformat() if core_properties.created else '',
            'modified': core_properties.modified.isoformat() if core_properties.modified else '',
            'paragraph_count': len(doc.paragraphs),
            'word_count': len(content.split())
        }

        return {
            'content': content,
            'metadata': metadata
        }
    except Exception as e:
        logger.error(f"Error processing DOCX file {file_name}: {str(e)}")
        raise

def process_csv_file(file_path: str, file_name: str) -> Dict[str, Any]:
    """Process a CSV file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            # Try to detect the dialect
            sample = file.read(1024)
            file.seek(0)

            sniffer = csv.Sniffer()
            dialect = sniffer.sniff(sample)
            has_header = sniffer.has_header(sample)

            # Read the CSV file
            reader = csv.reader(file, dialect)
            rows = list(reader)

            # Extract headers if present
            headers = []
            if has_header and rows:
                headers = rows[0]
                rows = rows[1:]

            # Convert to text content
            if has_header:
                content = ", ".join(headers) + "\n"
            else:
                content = ""

            for row in rows:
                content += ", ".join(row) + "\n"

            # Extract metadata
            file_stats = os.stat(file_path)
            metadata = {
                'source': file_name,
                'file_type': 'csv',
                'file_size': file_stats.st_size,
                'row_count': len(rows),
                'column_count': len(headers) if has_header else (len(rows[0]) if rows else 0),
                'has_header': has_header,
                'headers': headers if has_header else []
            }

            return {
                'content': content,
                'metadata': metadata
            }
    except Exception as e:
        logger.error(f"Error processing CSV file {file_name}: {str(e)}")
        raise

# -----------------------------------------------------------------------------
# Embedding Models
# -----------------------------------------------------------------------------

class SimpleEmbedding:
    """A simple embedding model with fallback to hash-based embeddings."""

    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        """
        Initialize the embedding model.

        Args:
            model_name: Name of the sentence transformer model to use
        """
        self.model_name = model_name
        self.dimension = 384  # Default dimension for fallback

        # Try to load the model if sentence-transformers is available
        if HAVE_SENTENCE_TRANSFORMERS:
            try:
                self.model = SentenceTransformer(model_name)
                self.dimension = self.model.get_sentence_embedding_dimension()
                logger.info(f"Loaded SentenceTransformer model: {model_name}")
            except Exception as e:
                logger.error(f"Error loading SentenceTransformer model: {e}")
                self.model = None
        else:
            self.model = None
            logger.warning("SentenceTransformer not available, using fallback embeddings")

    def _hash_embedding(self, text: str) -> List[float]:
        """Generate a hash-based embedding for fallback."""
        # Initialize a vector of zeros
        vector = [0.0] * self.dimension

        # Use words to influence different dimensions
        words = text.lower().split()
        for i, word in enumerate(words):
            # Hash the word
            hash_value = int(hashlib.md5(word.encode()).hexdigest(), 16)

            # Use the hash to set values in the vector
            for j in range(min(10, len(word))):
                idx = (hash_value + j) % self.dimension
                vector[idx] = 0.1 * ((hash_value % 20) - 10) + vector[idx]

        # Normalize the vector
        return self._normalize_vector(vector)

    def _normalize_vector(self, vector: List[float]) -> List[float]:
        """Normalize a vector to unit length."""
        # Convert to numpy array for efficient operations
        np_vector = np.array(vector, dtype=np.float32)
        norm = np.linalg.norm(np_vector)

        # Avoid division by zero
        if norm > 0:
            normalized = np_vector / norm
            return normalized.tolist()
        return vector

    def embed_text(self, text: str) -> List[float]:
        """
        Generate embedding for a single text.

        Args:
            text: The text to embed

        Returns:
            List of floats representing the embedding vector
        """
        if self.model is not None:
            try:
                embedding = self.model.encode(text, convert_to_numpy=True)
                return embedding.tolist()
            except Exception as e:
                logger.warning(f"Error generating embedding with SentenceTransformer: {e}")
                return self._hash_embedding(text)
        else:
            return self._hash_embedding(text)

    def embed_documents(self, documents: List[str]) -> List[List[float]]:
        """
        Generate embeddings for multiple documents.

        Args:
            documents: List of texts to embed

        Returns:
            List of embedding vectors
        """
        if self.model is not None:
            try:
                embeddings = self.model.encode(documents, convert_to_numpy=True)
                return embeddings.tolist()
            except Exception as e:
                logger.warning(f"Error generating embeddings with SentenceTransformer: {e}")
                return [self._hash_embedding(doc) for doc in documents]
        else:
            return [self._hash_embedding(doc) for doc in documents]

# -----------------------------------------------------------------------------
# LLM Clients
# -----------------------------------------------------------------------------

class GroqClient:
    """Client for Groq API."""

    def __init__(self, api_key: Optional[str] = None, model: str = "llama3-8b-8192"):
        """
        Initialize the Groq client.

        Args:
            api_key: Groq API key
            model: Model to use for generation
        """
        self.model = model

        # Set API key from args or environment
        if api_key is None:
            self.api_key = os.environ.get("GROQ_API_KEY")
        else:
            self.api_key = api_key

        # Check if we can make API calls
        self.can_use_api = HAVE_REQUESTS and self.api_key is not None
        if not self.can_use_api:
            logger.warning("Groq API unavailable. API key missing or requests not installed.")

    def generate_text(self, prompt: str) -> str:
        """
        Generate text using the Groq API.

        Args:
            prompt: The prompt to generate from

        Returns:
            Generated text
        """
        if not self.can_use_api:
            return f"[Simulated response for: {prompt[:50]}...]"

        try:
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {self.api_key}"
            }

            payload = {
                "model": self.model,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.7,
                "max_tokens": 1024
            }

            response = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers=headers,
                json=payload
            )

            if response.status_code == 200:
                return response.json()
            else:
                logger.warning(f"API error: {response.status_code} - {response.text}")
                return f"[Error: API returned status code {response.status_code}]"

        except Exception as e:
            logger.warning(f"Error calling Groq API: {e}")
            return f"[Error: {str(e)}]"

    def extract_text_from_response(self, response: Any) -> str:
        """
        Extract text from API response.

        Args:
            response: API response

        Returns:
            Extracted text
        """
        if isinstance(response, str):
            return response

        try:
            if isinstance(response, dict):
                if "choices" in response:
                    return response["choices"][0]["message"]["content"]
                elif "error" in response:
                    return f"[Error: {response['error']['message']}]"
            return str(response)
        except Exception as e:
            logger.warning(f"Error extracting text from response: {e}")
            return str(response)

class SimpleLLMClient:
    """Simple LLM client that returns simulated responses."""

    def generate_text(self, prompt: str) -> str:
        """
        Generate simulated text response.

        Args:
            prompt: The prompt to generate from

        Returns:
            Simulated response
        """
        # Extract question from prompt
        question_match = re.search(r"Question: (.*?)(\n|$)", prompt)
        context_match = re.search(r"Context:(.*?)(\n\n|$)", prompt, re.DOTALL)

        question = question_match.group(1) if question_match else "unknown question"
        has_context = bool(context_match)

        # Generate a simple response based on the prompt
        if "summarize" in prompt.lower() or "summary" in prompt.lower():
            return f"This document discusses various topics related to {question}. It covers key points and provides detailed information about the subject matter."

        elif "synthesize" in prompt.lower():
            return f"Based on the provided sources, {question} involves multiple aspects. The information suggests that this topic has several important dimensions worth considering."

        elif has_context:
            return f"Based on the provided context, the answer to '{question}' appears to be related to the information in the documents. The documents contain relevant details that address this question."

        else:
            return f"I don't have enough information to provide a specific answer about '{question}'. Please provide more context or try a different question."

    def extract_text_from_response(self, response: str) -> str:
        """
        Extract text from response.

        Args:
            response: Response text

        Returns:
            Extracted text
        """
        # For the simple client, just return the response as is
        return response

# -----------------------------------------------------------------------------
# RAG System
# -----------------------------------------------------------------------------

class SimpleRAGSystem:
    """A simple RAG system with vector database integration."""

    def __init__(self, documents, embeddings, vector_store_type="faiss"):
        """
        Initialize the RAG system.

        Args:
            documents: List of document chunks
            embeddings: List of embedding vectors for the chunks
            vector_store_type: Type of vector database to use ("faiss" or "chroma")
        """
        self.documents = documents
        self.vector_store_type = vector_store_type.lower()

        # Initialize vector database
        if self.vector_store_type == "faiss":
            self.vector_db = self._init_faiss(documents, embeddings)
        elif self.vector_store_type == "chroma":
            self.vector_db = self._init_chroma(documents, embeddings)
        else:
            raise ValueError(f"Unsupported vector store type: {vector_store_type}")

    def _init_faiss(self, documents, embeddings):
        """Initialize a FAISS vector database."""
        if not FAISS_AVAILABLE:
            raise ImportError("FAISS is not available. Install with 'pip install faiss-cpu'")

        # Convert embeddings to numpy array
        embeddings_array = np.array(embeddings).astype('float32')

        # Create FAISS index
        dimension = embeddings_array.shape[1]
        index = faiss.IndexFlatL2(dimension)
        index.add(embeddings_array)

        return {
            "index": index,
            "documents": documents,
            "embeddings": embeddings_array
        }

    def _init_chroma(self, documents, embeddings):
        """Initialize a ChromaDB vector database."""
        if not CHROMA_AVAILABLE:
            raise ImportError("ChromaDB is not available. Install with 'pip install chromadb'")

        # Create ChromaDB client
        client = chromadb.Client()
        collection = client.create_collection("documents")

        # Add documents to collection
        for i, (doc, embedding) in enumerate(zip(documents, embeddings)):
            collection.add(
                ids=[f"doc_{i}"],
                embeddings=[embedding],
                metadatas=[doc.get("metadata", {})]
            )

        return {
            "collection": collection,
            "documents": documents
        }

    def _faiss_search(self, query_embedding, top_k=5):
        """Search using FAISS."""
        # Convert query embedding to numpy array
        query_embedding = np.array([query_embedding]).astype('float32')

        # Search
        distances, indices = self.vector_db["index"].search(query_embedding, top_k)

        # Get results
        results = []
        for i, (dist, idx) in enumerate(zip(distances[0], indices[0])):
            if idx < len(self.vector_db["documents"]):
                doc = self.vector_db["documents"][idx]
                results.append({
                    "content": doc.get("content", ""),
                    "metadata": doc.get("metadata", {}),
                    "score": float(1.0 / (1.0 + dist))  # Convert distance to similarity score
                })

        return results

    def _chroma_search(self, query_embedding, top_k=5):
        """Search using ChromaDB."""
        # Search
        results = self.vector_db["collection"].query(
            query_embeddings=[query_embedding],
            n_results=top_k
        )

        # Get results
        output = []
        for i, (id, distance, metadata) in enumerate(zip(
            results["ids"][0], results["distances"][0], results["metadatas"][0]
        )):
            # Get the original document
            idx = int(id.split("_")[1])
            doc = self.vector_db["documents"][idx]

            output.append({
                "content": doc.get("content", ""),
                "metadata": metadata,
                "score": float(1.0 / (1.0 + distance))  # Convert distance to similarity score
            })

        return output

    def _keyword_search(self, query, top_k=5):
        """Simple keyword-based search."""
        # Tokenize query
        query_terms = set(query.lower().split())

        # Calculate scores based on term frequency
        results = []
        for i, doc in enumerate(self.documents):
            content = doc.get("content", "").lower()
            score = sum(1 for term in query_terms if term in content)

            if score > 0:
                results.append({
                    "content": doc.get("content", ""),
                    "metadata": doc.get("metadata", {}),
                    "score": score / len(query_terms)  # Normalize score
                })

        # Sort by score and limit to top_k
        results.sort(key=lambda x: x["score"], reverse=True)
        return results[:top_k]

    def retrieve(self, query, embedding_model, top_k=5, use_hybrid=False):
        """
        Retrieve relevant documents for a query.

        Args:
            query: User question
            embedding_model: Model to generate query embedding
            top_k: Number of documents to retrieve
            use_hybrid: Whether to use hybrid retrieval (semantic + keyword)

        Returns:
            List of relevant document chunks
        """
        # Generate query embedding
        query_embedding = embedding_model.embed_text(query)

        # Semantic search
        if self.vector_store_type == "faiss":
            results = self._faiss_search(query_embedding, top_k)
        else:
            results = self._chroma_search(query_embedding, top_k)

        # Hybrid search (combine with keyword search)
        if use_hybrid:
            keyword_results = self._keyword_search(query, top_k)

            # Merge results (simple approach)
            all_results = {}

            # Add semantic results
            for result in results:
                content = result["content"]
                all_results[content] = result

            # Add keyword results with lower weight
            for result in keyword_results:
                content = result["content"]
                if content in all_results:
                    # Boost existing score
                    all_results[content]["score"] = (all_results[content]["score"] + 0.5 * result["score"]) / 1.5
                else:
                    # Add new result with reduced score
                    result["score"] *= 0.7
                    all_results[content] = result

            # Convert back to list and sort
            results = list(all_results.values())
            results.sort(key=lambda x: x["score"], reverse=True)

        return results[:top_k]

    def analyze_question(self, question):
        """
        Analyze a question to identify its type and key entities.

        Args:
            question: User question

        Returns:
            Dictionary with question analysis
        """
        # Identify question type
        question_types = {
            "what": "definition",
            "who": "person",
            "when": "time",
            "where": "location",
            "why": "reason",
            "how": "process"
        }

        question_lower = question.lower()
        question_type = "general"

        for q_word, q_type in question_types.items():
            if question_lower.startswith(q_word):
                question_type = q_type
                break

        # Extract key entities (simplified)
        words = question_lower.split()
        stop_words = {"what", "who", "when", "where", "why", "how", "is", "are", "the", "a", "an"}
        entities = [word for word in words if word not in stop_words and len(word) > 3]

        # Check if it's a metadata query
        is_metadata_query = any(term in question_lower for term in
                               ["metadata", "author", "date", "title", "file type", "created",
                                "modified", "how many", "count", "written by"])

        return {
            "type": question_type,
            "entities": entities,
            "is_metadata_query": is_metadata_query
        }

    def retrieve_metadata(self, question):
        """
        Retrieve documents based on metadata query.

        Args:
            question: User question

        Returns:
            List of relevant documents
        """
        question_lower = question.lower()
        results = []

        # Extract potential metadata fields from question
        metadata_fields = {
            "author": ["author", "written by", "creator"],
            "title": ["title", "named", "called"],
            "date": ["date", "when", "created on", "modified on"],
            "file_type": ["file type", "format", "pdf", "docx", "txt", "csv"]
        }

        # Identify which fields are being queried
        target_fields = []
        for field, keywords in metadata_fields.items():
            if any(keyword in question_lower for keyword in keywords):
                target_fields.append(field)

        # If no specific fields identified, return all metadata
        if not target_fields:
            target_fields = list(metadata_fields.keys())

        # Extract entities that might be values
        words = question_lower.split()
        stop_words = {"what", "who", "when", "where", "why", "how", "is", "are", "the", "a", "an",
                     "metadata", "document", "file", "about", "contains", "have"}
        entities = [word for word in words if word not in stop_words and len(word) > 3]

        # Score documents based on metadata matches
        for doc in self.documents:
            metadata = doc.get("metadata", {})
            score = 0

            # Check for matches in target fields
            for field in target_fields:
                field_value = str(metadata.get(field, "")).lower()

                # Direct match with field value
                if field_value:
                    for entity in entities:
                        if entity in field_value:
                            score += 2

                    # Exact match bonus
                    if any(entity == field_value for entity in entities):
                        score += 3

            # Add document if it has a score
            if score > 0:
                results.append({
                    "content": doc.get("content", ""),
                    "metadata": metadata,
                    "score": score
                })

        # Sort by score
        results.sort(key=lambda x: x["score"], reverse=True)
        return results

    def answer_metadata_query(self, question, metadata_results, llm_client):
        """
        Answer a metadata query.

        Args:
            question: User question
            metadata_results: Metadata query results
            llm_client: LLM client for text generation

        Returns:
            Generated answer
        """
        if not metadata_results:
            return "I couldn't find any documents with metadata matching your query."

        # Format metadata for display
        metadata_text = ""
        for i, result in enumerate(metadata_results[:5]):
            metadata = result.get("metadata", {})
            metadata_text += f"Document {i+1}:\n"

            for key, value in metadata.items():
                metadata_text += f"  - {key}: {value}\n"

            metadata_text += "\n"

        # Create prompt for the LLM
        prompt = f"""
        Answer the following question about document metadata based on the provided information.

        Question: {question}

        Document Metadata:
        {metadata_text}

        Answer:
        """

        # Generate answer
        response = llm_client.generate_text(prompt)
        answer = llm_client.extract_text_from_response(response)

        return answer

    def assess_confidence(self, question, chunks):
        """
        Assess confidence in retrieved chunks.

        Args:
            question: User question
            chunks: Retrieved chunks

        Returns:
            Confidence score (0-1)
        """
        if not chunks:
            return 0.0

        # Extract question entities
        question_lower = question.lower()
        words = question_lower.split()
        stop_words = {"what", "who", "when", "where", "why", "how", "is", "are", "the", "a", "an"}
        question_entities = set(word for word in words if word not in stop_words and len(word) > 3)

        # Calculate entity coverage
        total_coverage = 0
        for chunk in chunks:
            content = chunk.get("content", "").lower()
            chunk_score = chunk.get("score", 0)

            # Count how many question entities appear in the chunk
            covered_entities = sum(1 for entity in question_entities if entity in content)
            entity_coverage = covered_entities / len(question_entities) if question_entities else 0

            # Combine with retrieval score
            chunk_confidence = (0.7 * entity_coverage + 0.3 * chunk_score)
            total_coverage += chunk_confidence

        # Average confidence across chunks, with diminishing returns for more chunks
        confidence = total_coverage / (len(chunks) ** 0.8)

        # Normalize to 0-1 range
        return min(max(confidence, 0.0), 1.0)

    def generate_answer_with_citations(self, question, chunks, llm_client):
        """
        Generate an answer with citations.

        Args:
            question: User question
            chunks: Retrieved chunks
            llm_client: LLM client for text generation

        Returns:
            Answer with citations
        """
        if not chunks:
            return "I couldn't find any relevant information to answer your question."

        # Prepare context from chunks with source identifiers
        contexts = []
        for i, chunk in enumerate(chunks):
            source = chunk.get("metadata", {}).get("source", f"Source {i+1}")
            contexts.append(f"[{i+1}] {chunk['content']} (Source: {source})")

        context_text = "\n\n".join(contexts)

        # Create prompt for the LLM
        prompt = f"""
        Answer the following question based on the provided context.
        Include citations [1], [2], etc. to indicate which sources support your answer.
        If the context doesn't contain relevant information, say "I don't have enough information to answer this question."

        Context:
        {context_text}

        Question: {question}

        Answer (with citations):
        """

        # Generate answer
        response = llm_client.generate_text(prompt)
        answer = llm_client.extract_text_from_response(response)

        return answer

    def generate_response_with_uncertainty(self, question, chunks, confidence, llm_client):
        """
        Generate a response with uncertainty handling.

        Args:
            question: User question
            chunks: Retrieved chunks
            confidence: Confidence score
            llm_client: LLM client for text generation

        Returns:
            Response with uncertainty acknowledgment
        """
        if not chunks:
            return "I don't have enough information to answer this question."

        # Prepare context from chunks
        context = "\n\n".join([chunk["content"] for chunk in chunks])

        # Create prompt with uncertainty guidance based on confidence
        if confidence < 0.3:
            uncertainty_guidance = "Express high uncertainty in your answer. Make it clear that you're not confident and the information might not be reliable."
        elif confidence < 0.7:
            uncertainty_guidance = "Express moderate uncertainty in your answer. Acknowledge that you're not entirely confident."
        else:
            uncertainty_guidance = "You can be relatively confident in your answer, but still acknowledge any limitations in the information."

        prompt = f"""
        Answer the following question based on the provided context.
        {uncertainty_guidance}

        Context:
        {context}

        Question: {question}

        Answer:
        """

        # Generate answer
        response = llm_client.generate_text(prompt)
        answer = llm_client.extract_text_from_response(response)

        return answer

    def synthesize_information(self, question, chunks, llm_client):
        """
        Synthesize information from multiple chunks.

        Args:
            question: User question
            chunks: Retrieved chunks
            llm_client: LLM client for text generation

        Returns:
            Synthesized information
        """
        if not chunks:
            return "I couldn't find any relevant information to synthesize."

        # Prepare context from chunks with source information
        contexts = []
        for i, chunk in enumerate(chunks):
            source = chunk.get("metadata", {}).get("source", f"Source {i+1}")
            contexts.append(f"Source {i+1} ({source}): {chunk['content']}")

        context_text = "\n\n".join(contexts)

        # Create synthesis prompt
        prompt = f"""
        Synthesize information from the following sources to answer the question.
        If the sources contain conflicting information, acknowledge the differences.

        Question: {question}

        Sources:
        {context_text}

        Synthesized Answer:
        """

        # Generate synthesized answer
        response = llm_client.generate_text(prompt)
        answer = llm_client.extract_text_from_response(response)

        return answer

# -----------------------------------------------------------------------------
# Document QA System
# -----------------------------------------------------------------------------

class DocumentQASystem:
    """Complete Document Q&A system."""

    def __init__(self, rag_system, embedding_model=None, llm_client=None):
        """
        Initialize the Document Q&A system.

        Args:
            rag_system: RAG system for retrieval
            embedding_model: Model for generating embeddings
            llm_client: Client for LLM text generation
        """
        self.rag_system = rag_system
        self.embedding_model = embedding_model

        # Initialize LLM client
        if llm_client is None:
            if HAVE_REQUESTS:
                self.llm_client = GroqClient()
                logger.info("Using GroqClient for text generation.")
            else:
                self.llm_client = SimpleLLMClient()
                logger.info("Using SimpleLLMClient for text generation (simulated responses).")
        else:
            self.llm_client = llm_client

    def answer_question(self, question, k=5, use_hybrid=False):
        """
        Answer a user question.

        Args:
            question: User question
            k: Number of chunks to retrieve
            use_hybrid: Whether to use hybrid retrieval

        Returns:
            Answer with sources and confidence information
        """
        # Analyze the question
        analysis = self.rag_system.analyze_question(question)

        # Handle metadata queries differently
        if analysis["is_metadata_query"]:
            metadata_results = self.rag_system.retrieve_metadata(question)
            answer = self.rag_system.answer_metadata_query(
                question, metadata_results, self.llm_client
            )
            return {
                "answer": answer,
                "sources": [r["metadata"] for r in metadata_results[:3]],
                "is_metadata_query": True,
                "confidence": 0.9 if metadata_results else 0.1
            }

        # Retrieve relevant chunks
        chunks = self.rag_system.retrieve(
            question, self.embedding_model, top_k=k, use_hybrid=use_hybrid
        )

        # Assess confidence
        confidence = self.rag_system.assess_confidence(question, chunks)

        # Generate answer with citations
        if confidence >= 0.5:
            answer = self.rag_system.generate_answer_with_citations(
                question, chunks, self.llm_client
            )
        else:
            # For low confidence, use uncertainty handling
            answer = self.rag_system.generate_response_with_uncertainty(
                question, chunks, confidence, self.llm_client
            )

        return {
            "answer": answer,
            "sources": [chunk.get("metadata", {}) for chunk in chunks],
            "is_metadata_query": False,
            "confidence": confidence
        }

    def answer_with_synthesis(self, question, k=5):
        """
        Answer a question with information synthesis from multiple sources.

        Args:
            question: User question
            k: Number of chunks to retrieve

        Returns:
            Synthesized answer with sources
        """
        # Retrieve relevant chunks
        chunks = self.rag_system.retrieve(
            question, self.embedding_model, top_k=k
        )

        # Synthesize information
        answer = self.rag_system.synthesize_information(
            question, chunks, self.llm_client
        )

        return {
            "answer": answer,
            "sources": [chunk.get("metadata", {}) for chunk in chunks],
            "is_synthesized": True,
            "confidence": 0.7  # Fixed confidence for synthesis
        }

    def get_document_summary(self, document_id=None, source=None):
        """
        Generate a summary of a document.

        Args:
            document_id: ID of the document to summarize
            source: Source name to summarize

        Returns:
            Document summary
        """
        # Find the document
        if document_id is not None:
            if document_id < 0 or document_id >= len(self.rag_system.documents):
                return {"error": f"Document ID {document_id} not found"}

            document = self.rag_system.documents[document_id]
            chunks = [document]
        elif source is not None:
            # Find all chunks from the source
            chunks = []
            for doc in self.rag_system.documents:
                if doc.get("metadata", {}).get("source") == source:
                    chunks.append(doc)

            if not chunks:
                return {"error": f"No documents found from source '{source}'"}
        else:
            return {"error": "Either document_id or source must be provided"}

        # Create a summary prompt
        content = "\n\n".join([chunk.get("content", "") for chunk in chunks])

        prompt = f"""
        Provide a concise summary of the following document content:

        {content}

        Summary:
        """

        # Generate summary
        response = self.llm_client.generate_text(prompt)
        summary = self.llm_client.extract_text_from_response(response)

        return {
            "summary": summary,
            "document_count": len(chunks),
            "metadata": chunks[0].get("metadata", {})
        }

# Example usage
if __name__ == "__main__":
    # This code will run when the module is executed directly
    print("RAG Components module loaded successfully.")
